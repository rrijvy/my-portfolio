#!/usr/bin/env python
# Generates an updated CV .docx for Ashikur Rahman Rijvy (rrijvy)
# Layout follows a clean "title + bullet points" pattern.
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x49, 0x7D)
DARK = RGBColor(0x22, 0x22, 0x22)
GREY = RGBColor(0x66, 0x66, 0x66)

doc = Document()

# ---------- base style ----------
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10.5)
normal.font.color.rgb = DARK
normal.paragraph_format.space_after = Pt(3)
normal.paragraph_format.line_spacing = 1.04

# narrow margins
for s in doc.sections:
    s.top_margin = Inches(0.45)
    s.bottom_margin = Inches(0.45)
    s.left_margin = Inches(0.7)
    s.right_margin = Inches(0.7)


def set_bottom_border(p, color='1F497D', sz=8):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(sz))
    bottom.set(qn('w:space'), '2')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def heading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text.upper())
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = NAVY
    set_bottom_border(p)
    return p


def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    for run in p.runs:
        run.font.size = Pt(10.5)
    return p


def add_hyperlink(paragraph, url, text, color=NAVY, bold=False, size=None):
    part = paragraph.part
    rpr = OxmlElement('w:rPr')
    if bold:
        b = OxmlElement('w:b')
        rpr.append(b)
    if size is not None:
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(size * 2)))
        rpr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), str(int(size * 2)))
        rpr.append(szCs)
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '%02X%02X%02X' % (color[0], color[1], color[2]))
    rpr.append(c)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rpr.append(u)
    base = 'rId'
    existing = set(part.rels.keys())
    n = 1000
    while f'{base}{n}' in existing:
        n += 1
    rel = part.rels.add_relationship(
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        url, is_external=True, rId=f'{base}{n}')
    h = OxmlElement('w:hyperlink')
    h.set(qn('r:id'), rel.rId)
    run = OxmlElement('w:r')
    run.append(rpr)
    t = OxmlElement('w:t')
    t.text = text
    run.append(t)
    h.append(run)
    paragraph._p.append(h)
    return h


def proj_bullet(desc, name, url=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    if url:
        add_hyperlink(p, url, name, bold=True)
    else:
        r = p.add_run(name)
        r.bold = True
    p.add_run(': ' + desc)
    for run in p.runs:
        run.font.size = Pt(10.5)
    return p


def role(company, period, title, bullets):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(7.0), WD_TAB_ALIGNMENT.RIGHT)
    r = p.add_run(company)
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = DARK
    rp = p.add_run('\t' + period)
    rp.font.size = Pt(10); rp.italic = True
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(3)
    r2 = p2.add_run(title)
    r2.italic = True; r2.font.size = Pt(10.5); r2.font.color.rgb = GREY
    for b in bullets:
        if isinstance(b, tuple):
            bullet(b[1], bold_lead=b[0])
        else:
            bullet(b)


def exp_role(title, period, company, bullets, note=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(7.0), WD_TAB_ALIGNMENT.RIGHT)
    r = p.add_run(title)
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = DARK
    rp = p.add_run('\t' + period)
    rp.font.size = Pt(10); rp.italic = True
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(1) if note else Pt(3)
    r2 = p2.add_run(company)
    r2.bold = True; r2.font.size = Pt(10.5); r2.font.color.rgb = NAVY
    if note:
        p3 = doc.add_paragraph()
        p3.paragraph_format.space_after = Pt(3)
        r3 = p3.add_run(note)
        r3.italic = True; r3.font.size = Pt(9.5); r3.font.color.rgb = GREY
    for b in bullets:
        bullet(b)


def notable_block(title, context, role_text, challenge, solution, outcome, tech):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(title)
    r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = DARK
    rc = p.add_run('   |   ' + context)
    rc.font.size = Pt(9.5); rc.italic = True; rc.font.color.rgb = GREY
    pr = doc.add_paragraph()
    pr.paragraph_format.space_after = Pt(1)
    rr = pr.add_run('Role: ')
    rr.bold = True; rr.font.size = Pt(10); rr.font.color.rgb = DARK
    rr2 = pr.add_run(role_text)
    rr2.font.size = Pt(10); rr2.font.color.rgb = DARK
    bullet(challenge, bold_lead='Challenge: ')
    bullet(solution, bold_lead='Solution: ')
    bullet(tech, bold_lead='Tech Stack: ')
    bullet(outcome, bold_lead='Outcome: ')


# ================= HEADER =================
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.paragraph_format.space_after = Pt(1)
r = t.add_run('ASHIKUR RAHMAN RIJVY')
r.bold = True; r.font.size = Pt(22); r.font.color.rgb = NAVY

ttl = doc.add_paragraph()
ttl.alignment = WD_ALIGN_PARAGRAPH.CENTER
ttl.paragraph_format.space_after = Pt(3)
rt = ttl.add_run('Senior Software Engineer')
rt.bold = True; rt.font.size = Pt(13); rt.font.color.rgb = DARK

c = doc.add_paragraph()
c.alignment = WD_ALIGN_PARAGRAPH.CENTER
c.paragraph_format.space_after = Pt(1)

def add_contact(text, url=None, sep=None, size=9):
    if url:
        add_hyperlink(c, url, text, color=RGBColor(0x1F, 0x49, 0x7D), size=size)
    else:
        r = c.add_run(text)
        r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    if sep:
        rs = c.add_run(sep)
        rs.font.size = Pt(9); rs.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

add_contact('Dhaka, Bangladesh', sep='  |  ')
add_contact('Mobile: +8801717745808', url='tel:+8801717745808', sep='  |  ')
add_contact('rrijvy@gmail.com', url='mailto:rrijvy@gmail.com')
c2 = doc.add_paragraph()
c2.alignment = WD_ALIGN_PARAGRAPH.CENTER
c2.paragraph_format.space_after = Pt(2)

def add_chip(label, url):
    add_hyperlink(c2, url, label, color=RGBColor(0x1F, 0x49, 0x7D), size=9)
    sep = c2.add_run('   |   ')
    sep.font.size = Pt(9); sep.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

add_chip('LinkedIn', 'https://linkedin.com/in/rrijvy')
add_chip('GitHub', 'https://github.com/rrijvy')
add_chip('Portfolio', 'https://rrijvy.github.io/my-portfolio')

# ================= SUMMARY =================
heading('Summary')
obj = ('Software engineer with 7+ years of experience building SaaS and document-automation products in .NET '
       'and TypeScript/Node, mostly on AWS and Azure. At KAZ Software I build serverless EDI pipelines, document '
       'extraction with LLMs, regulatory AI tools, and e-invoicing compliance systems. I like clean, simple code '
       'and reliable delivery, and I build and share open-source projects in my free time.')
p = doc.add_paragraph(obj)
p.paragraph_format.space_after = Pt(2)

# ================= TECHNICAL SKILLS =================
heading('Technical Skills')
bullet('C# / .NET (7 yrs), TypeScript (6 yrs), JavaScript, Python, Node.js', bold_lead='Languages: ')
bullet('React, Next.js, Vite, Tailwind CSS, shadcn-ui, Redux/Zustand/TanStack Query, Electron, Expo / React Native', bold_lead='Frontend: ')
bullet('.NET Core / .NET 10, ASP.NET MVC & Razor, Node/Express, FastAPI, REST APIs, JWT (RS256)', bold_lead='Backend: ')
bullet('AWS (5 yrs: Lambda, Step Functions, DynamoDB, S3, SNS, API Gateway, SES, EventBridge), Azure (Functions, Blob), Docker; pnpm/turbo monorepos; CI/CD', bold_lead='Cloud & Infra: ')
bullet('AWS Comprehend, LayoutLMv3, Claude / GPT-4 / Bedrock, ModernBERT, ComfyUI, Flux, LoRA training, Segment Anything (SAM), Hugging Face, Web Speech API', bold_lead='AI / ML: ')
bullet('PostgreSQL, MongoDB, MSSQL Server; Prisma, Entity Framework (ORM)', bold_lead='Data: ')
bullet('Git, multi-tenant & RBAC architecture, design patterns, clean/hexagonal architecture', bold_lead='Practices: ')

# ================= NOTABLE PROJECT ACHIEVEMENTS =================
heading('Notable Project Achievements')
notable_block('Document Processing & EDI Pipeline', 'KAZ Software',
    'Senior Software Engineer',
    'Supplier documents came in by email, usually in batches of about 25 every 1-2 hours, and someone had to process each one by hand. This was slow and led to data entry mistakes.',
    'I built a pipeline using AWS Step Functions that classifies each document with AWS Comprehend, and falls back to our own trained LayoutLMv3 model for harder cases. Claude, GPT-4 and AWS Bedrock then read the document and pull out the line-item data, which gets checked against our business rules.',
    'No more manual handling - the system now processes 500-1,000 documents a month on its own, at about 95% accuracy.',
    'AWS Comprehend, LayoutLMv3, Claude / GPT-4 / Bedrock, AWS Lambda / Step Functions')
notable_block('UAE E-Invoicing Compliance SaaS', 'KAZ Software',
    'Technical Lead & Project Manager',
    'Businesses had to check their invoices against the UAE PINT-AE standard by hand, which took 1-2 hours per invoice, with no easy way to catch problems before going live.',
    'I built a multi-tenant platform where businesses upload their invoice data as a CSV, and it gets checked against rules they define themselves in a rule builder. It also flags exceptions for review and puts together evidence for audits.',
    'Validation now takes 10-15 seconds per batch instead of 1-2 hours per invoice, and it flags about 25% of invoices for review before submission. It is live with 2-3 tenants in the UAE.',
    'React, Node / Express, PostgreSQL, Azure')
notable_block('Regulatory Intelligence Platform', 'KAZ Software',
    'Technical Lead & Project Manager',
    'An analyst was spending 5-6 hours a day manually checking 12 news and regulatory sources just to stay on top of risk.',
    'I built a multi-agent system with RAG that checks those sources automatically every day. Each site gets its own scraper since they are all built differently, and the system rotates browser fingerprints to stay reliable. It then puts together AI summaries and risk insights.',
    'This cut daily research time by about 87%, from 5-6 hours down to 40-45 minutes, across all 12 sources.',
    'Python, FastAPI, Celery, LangChain / LangGraph, Pinecone (RAG)')
notable_block('Transformer Risk Classifier', 'KAZ Software',
    'ML Engineer',
    'Tax and regulatory articles needed to be sorted by risk category, and doing this by hand was slow and inconsistent.',
    'I trained a ModernBERT classifier that covers 5 main categories and roughly 25-30 subcategories.',
    'It classifies each article in under 20 seconds at about 98% accuracy, handling around 50 articles a day.',
    'ModernBERT, FastAPI, Docker')
notable_block('Children’s Comic Book Illustration Engine', 'KAZ Software',
    'Generative AI Engineer',
    'A client needed illustrations for a children’s comic book, with the same custom characters staying consistent across every panel. There was no way to do this without a manual illustration process, and no way to edit just one part of an image at a time.',
    'I built the whole image-generation engine myself in ComfyUI, using the open-source Flux model with custom-trained LoRAs to keep characters consistent. I also built a custom ComfyUI node on top of Segment Anything (SAM) so I could pick just the part of an image that needed to change and regenerate only that part.',
    'It now generates a full set of 15 illustrations in about 20 minutes, with consistent characters throughout thanks to the trained LoRAs.',
    'ComfyUI, Flux (diffusion model), custom LoRA training, Segment Anything (SAM), Python')

# ================= PROFESSIONAL EXPERIENCE =================
heading('Professional Experience')

exp_role('Senior Software Engineer', 'Jul 2020 - Present', 'KAZ Software', [
    'Led the build of a multi-tenant B2B supply-chain platform that automates order processing, supplier and purchase-order management, and shipment tracking.',
    'Added in-app discussion threads with mentions and reactions, a file manager backed by S3 with previews for different file types, and a notification system where users can turn email and in-app alerts on or off separately for each type, plus a daily digest email.',
    'Built an Open API (documented with Swagger, secured with API keys) so third parties could integrate with the platform, and added a WebSocket connection so users see order and discussion updates right away.',
    'Built a serverless document-extraction pipeline (AWS Step Functions and Lambda) that classifies documents with AWS Comprehend, with an in-house LayoutLMv3 model as a fallback, then reads and extracts line-item data using Claude, GPT-4 and Bedrock. It now processes 500-1,000 supplier documents a month at about 95% accuracy.',
    'Built a UAE PINT-AE e-invoicing compliance product with a rule builder so compliance teams can write their own validation rules. It cut invoice validation time from 1-2 hours to 10-15 seconds per batch, and it is now live with 2-3 tenants.',
    'Built a multi-agent regulatory-intelligence system with RAG that automatically checks 12 news and regulatory sources every day, using a different scraper for each site and rotating browser fingerprints. This cut analyst research time by about 87%.',
    'Trained a ModernBERT classifier for tax and regulatory articles, covering 5 categories and roughly 25-30 subcategories, at about 98% accuracy on around 50 articles a day.',
    'Built a scheduled Electron app that connects to 8 different ERP systems, including Progress OpenEdge, Visual FoxPro, P21/Epicor and Plex, through direct database and REST connections, keeping purchase-order line items and documents in sync on a schedule.',
    'Built an image-generation engine in ComfyUI for a children’s comic-book client, using Flux and custom-trained LoRAs to keep characters consistent, plus a custom Segment Anything (SAM) node for editing just one part of an image. It generates 15 illustrations in about 20 minutes.',
], note='Promoted July 2024 - previously Software Engineer & Associate Software Engineer')

exp_role('Software Developer', 'May 2019 - Feb 2020', 'Alphasoft Technology Limited', [
    'Built a hospital management system covering patient assessment, billing, live dashboards, and real-time alerts between doctors and receptionists using SignalR.',
    'Built a website generator that let non-technical staff create a full website just by filling out a form, so nobody had to hand-code each site.',
    'Built a payroll system that calculates salaries from job cards and runs an approval workflow based on attendance.',
])

# ================= OPEN SOURCE / PERSONAL =================
heading('Open Source & Personal Projects')

GH = 'https://github.com/rrijvy/'
proj_bullet('Online store and admin console: shoppers browse and buy products, while staff manage catalog, orders and stock. Built as a full-stack app (Next.js 15 + .NET 10 + MongoDB, AWS Lambda).', 'Ecommerce Platform', GH + 'ecommerce')
proj_bullet('Operations hub for a small contracting and tender business: tracks project finances, working capital, assets, staff and deadline reminders so the owner sees profit/loss and dues in one place (React + Vite + Express + Prisma, AWS SAM).', 'Tender Operations Platform', GH + 'FFEnterprise')
proj_bullet('Voice typing for any website: a Chrome extension that turns speech into text in any input field, in English and Bengali, so users can write without typing (Manifest V3, Web Speech + Google Cloud).', 'Vocaly', GH + 'Vocaly')
proj_bullet('Embeddable PDF viewer for React apps, with selectable text, so developers can show documents without a heavy dependency (pdfjs-dist, npm).', 'RX PDF Viewer', GH + 'rx-pdf-viewer')
proj_bullet('Clipboard history manager: a Chrome extension that remembers recently copied text locally so users can re-copy it later, storing nothing externally (TypeScript).', 'Colorful Copy Manager', GH + 'colorful-copy-manager')
proj_bullet('Course-selling site: creators publish video courses with secure streaming, subscriptions and certificates, and learners track progress (Next.js + Express + Prisma).', 'EdTech Platform', GH + 'EdTech')
proj_bullet('Distributed web crawler: collects and scrapes web content at scale using worker queues and multiple datastores (.NET, Python/Django, Next.js, Postgres/Mongo/Redis, Celery).', 'Bebodh Crawler', GH + 'BebodhCrawler')
proj_bullet('Job-alert app: lets job hunters save search filters and get matched openings in one mobile feed, with application tracking (Expo/React Native).', 'UpAlert', GH + 'upalert')

# ================= EDUCATION =================
heading('Education')
role('Daffodil International University', '2019', 'BSc in Computer Science and Engineering (GPA 3.1 / 4.0)', [])
role('Narail Govt Victoria College', '2013', 'Higher Secondary School Certificate (GPA 4.80)', [])
role('Narail Govt High School', '2011', 'Secondary School Certificate (GPA 4.75)', [])

# ================= REFERENCES =================
heading('References')
ref = ('Shawal Siddique Shaon              Hannan Hossain\n'
       'CTO, KAZ Software                   Principal Software Engineer, KAZ Software\n'
       'Cell: 01617607518                   Cell: 01710487883')
rp = doc.add_paragraph(ref)
rp.paragraph_format.space_after = Pt(2)
for run in rp.runs:
    run.font.size = Pt(10.5)

OUT = r'D:\downloads\Resume_of_Rijvy_v7.docx'
doc.save(OUT)
print('Saved: ' + OUT)
